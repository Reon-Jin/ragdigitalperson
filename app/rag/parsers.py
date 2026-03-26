from __future__ import annotations

import logging
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import fitz
from docx import Document

from app.config import Settings
from app.rag.types import ExtractedDocument, PageContent


logger = logging.getLogger(__name__)


class DocumentParser:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def parse(self, doc_id: str, file_path: str | Path, filename: str, source_type: str = "upload") -> ExtractedDocument:
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            pages = self._parse_pdf(file_path)
        elif suffix == ".docx":
            pages = self._parse_docx(file_path)
        elif suffix in {".txt", ".md"}:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            pages = [PageContent(page_number=1, text=self._clean_text(text))]
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        text = "\n\n".join(page.text for page in pages if page.text.strip()).strip()
        return ExtractedDocument(doc_id=doc_id, filename=filename, suffix=suffix, text=text, pages=pages, source_type=source_type)

    def _parse_pdf(self, file_path: Path) -> list[PageContent]:
        document = fitz.open(file_path)
        try:
            page_indexes = list(range(document.page_count))
        finally:
            document.close()
        results: dict[int, PageContent] = {}
        with ThreadPoolExecutor(max_workers=max(1, self.settings.extraction_workers)) as executor:
            future_map = {executor.submit(self._extract_pdf_page, file_path, index): index for index in page_indexes}
            for future in as_completed(future_map):
                index = future_map[future]
                try:
                    results[index] = future.result()
                except Exception as exc:
                    logger.warning("pdf page extraction failed: file=%s page=%s error=%s", file_path, index + 1, exc)
                    results[index] = PageContent(page_number=index + 1, text="", metadata={"warning": str(exc)})
        return [results[index] for index in sorted(results)]

    def _extract_pdf_page(self, file_path: Path, page_index: int) -> PageContent:
        document = fitz.open(file_path)
        try:
            page = document.load_page(page_index)
            blocks: list[dict[str, object]] = []
            text_segments: list[str] = []
            raw = page.get_text("dict")
            for block in raw.get("blocks", []):
                lines: list[str] = []
                for line in block.get("lines", []):
                    spans = [span.get("text", "") for span in line.get("spans", []) if str(span.get("text", "")).strip()]
                    if spans:
                        lines.append("".join(spans).strip())
                if not lines:
                    continue
                joined = "\n".join(lines).strip()
                block_kind = "heading" if self._looks_like_heading(joined) else "text"
                if self._looks_like_table(joined):
                    block_kind = "table"
                blocks.append({"kind": block_kind, "text": joined, "bbox": block.get("bbox")})
                text_segments.append(joined)
            text = self._clean_text("\n\n".join(text_segments))
            return PageContent(page_number=page_index + 1, text=text, blocks=blocks)
        finally:
            document.close()

    def _parse_docx(self, file_path: Path) -> list[PageContent]:
        doc = Document(str(file_path))
        blocks: list[str] = []
        for paragraph in doc.paragraphs:
            text = self._clean_text(paragraph.text)
            if not text:
                continue
            if paragraph.style and "heading" in paragraph.style.name.lower():
                blocks.append(f"# {text}")
            else:
                blocks.append(text)
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                values = [self._clean_text(cell.text) for cell in row.cells]
                values = [value for value in values if value]
                if values:
                    rows.append(" | ".join(values))
            if rows:
                blocks.append("\n".join(rows))
        return [PageContent(page_number=1, text="\n\n".join(blocks).strip())]

    def _clean_text(self, text: str) -> str:
        cleaned = (text or "").replace("\r", "\n").replace("\xa0", " ")
        cleaned = re.sub(r"[ \t]+", " ", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def _looks_like_heading(self, text: str) -> bool:
        probe = text.strip()
        if len(probe) < 3 or len(probe) > 80:
            return False
        if re.match(r"^(\d+(\.\d+){0,3}|第[一二三四五六七八九十百千]+[章节部分])", probe):
            return True
        return probe == probe.upper() or (probe.endswith(":") and len(probe) <= 40)

    def _looks_like_table(self, text: str) -> bool:
        return "|" in text or "\t" in text or len(re.findall(r"\d+(?:\.\d+)?%?", text)) >= 3
